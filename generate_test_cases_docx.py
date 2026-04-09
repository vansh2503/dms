"""
Generate DMS_Test_Cases.docx – Hyundai Dealer Management System
Run: python generate_test_cases_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ──────────────────────────────────────────────
# Colour palette
# ──────────────────────────────────────────────
HYUNDAI_BLUE   = RGBColor(0x00, 0x2C, 0x5F)   # #002C5F  dark navy
ACCENT_BLUE    = RGBColor(0x00, 0x6A, 0xC6)   # #006AC6  medium blue
HEADER_BG      = RGBColor(0x00, 0x2C, 0x5F)   # table header fill
ALT_ROW        = RGBColor(0xE8, 0xF0, 0xFE)   # #E8F0FE  alternate row tint
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT      = RGBColor(0x1A, 0x1A, 0x2E)

# ──────────────────────────────────────────────
# Priority badge colours
# ──────────────────────────────────────────────
PRIORITY_COLORS = {
    "High":   RGBColor(0xC0, 0x39, 0x2B),
    "Medium": RGBColor(0xE6, 0x7E, 0x22),
    "Low":    RGBColor(0x27, 0xAE, 0x60),
}

TYPE_COLORS = {
    "Functional": RGBColor(0x15, 0x65, 0xC0),
    "Negative":   RGBColor(0xB7, 0x1C, 0x1C),
    "Security":   RGBColor(0x6A, 0x1B, 0x9A),
    "Boundary":   RGBColor(0xE6, 0x5C, 0x00),
}

# ──────────────────────────────────────────────
# Helper: set cell background colour
# ──────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, border_color="002C5F"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ──────────────────────────────────────────────
# Helper: add a styled paragraph inside a cell
# ──────────────────────────────────────────────
def cell_para(cell, text, bold=False, color=None, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else DARK_TEXT
    return run

# ──────────────────────────────────────────────
# Helper: header row for a section table
# ──────────────────────────────────────────────
COLUMNS = [
    ("TC ID",          Cm(1.8)),
    ("Module",         Cm(2.3)),
    ("Test Case Title",Cm(4.5)),
    ("Preconditions",  Cm(4.5)),
    ("Test Steps",     Cm(5.5)),
    ("Expected Result",Cm(5.0)),
    ("Test Type",      Cm(2.4)),
    ("Priority",       Cm(1.7)),
]

def add_table_header(table):
    hdr = table.rows[0]
    for i, (label, _) in enumerate(COLUMNS):
        cell = hdr.cells[i]
        set_cell_bg(cell, HEADER_BG)
        set_cell_border(cell, "FFFFFF")
        cell_para(cell, label, bold=True, color=WHITE, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

def set_col_widths(table):
    for i, (_, w) in enumerate(COLUMNS):
        for cell in table.column_cells(i):
            cell.width = w

# ──────────────────────────────────────────────
# Test case data
# ──────────────────────────────────────────────
TEST_CASES = [
    # ── Authentication ──
    {
        "id": "TC_001", "module": "Authentication",
        "title": "Successful login with valid credentials",
        "pre": "Backend running; user admin@hyundai.com exists with role SUPER_ADMIN.",
        "steps": "1. Navigate to /login.\n2. Enter email admin@hyundai.com and password admin123.\n3. Click Login.",
        "expected": "HTTP 200; JWT token stored in localStorage; user redirected to Dashboard; role shown as SUPER_ADMIN.",
        "type": "Functional", "priority": "High",
    },
    {
        "id": "TC_002", "module": "Authentication",
        "title": "Login fails with invalid password",
        "pre": "Backend running; user sales@hyundai.com exists.",
        "steps": "1. Navigate to /login.\n2. Enter email sales@hyundai.com and password wrongPass.\n3. Click Login.",
        "expected": "HTTP 401; error message 'Invalid credentials' shown; user stays on Login page; no token stored.",
        "type": "Negative", "priority": "High",
    },
    {
        "id": "TC_003", "module": "Authentication",
        "title": "Login with empty email and password",
        "pre": "Login page is accessible.",
        "steps": "1. Navigate to /login.\n2. Leave both fields blank.\n3. Click Login.",
        "expected": "Form validation triggers; fields marked required; API call NOT made; no token stored.",
        "type": "Negative", "priority": "High",
    },
    {
        "id": "TC_004", "module": "Authentication",
        "title": "API access without JWT token",
        "pre": "Backend running.",
        "steps": "1. Send GET /api/customers without Authorization header.",
        "expected": "HTTP 401 Unauthorized; response body contains error indicating missing/invalid token.",
        "type": "Security", "priority": "High",
    },
    {
        "id": "TC_005", "module": "Authentication",
        "title": "JWT token expires and session ends",
        "pre": "Valid user is logged in; token TTL is configured.",
        "steps": "1. Log in as a valid user.\n2. Wait until token expires (or manually alter exp claim).\n3. Navigate to /customers.",
        "expected": "User redirected to Login page; stored token cleared; HTTP 401 returned on any API call.",
        "type": "Security", "priority": "High",
    },
    # ── Customer ──
    {
        "id": "TC_006", "module": "Customer",
        "title": "Add new customer with all valid fields",
        "pre": "Logged in as DEALER_MANAGER or SALES_EXECUTIVE.",
        "steps": "1. Navigate to Customers → Add Customer.\n2. Fill first name, last name, phone, email, city, customer type (RETAIL).\n3. Click Save.",
        "expected": "HTTP 201; new customer record appears in list with correct details; no duplicate created.",
        "type": "Functional", "priority": "High",
    },
    {
        "id": "TC_007", "module": "Customer",
        "title": "Search customer by name keyword",
        "pre": "At least one customer with first name 'Rahul' exists.",
        "steps": "1. Navigate to Customers.\n2. Enter 'Rahul' in the search bar.\n3. Observe filtered results.",
        "expected": "Only customers whose name contains 'Rahul' are displayed; empty state shown if no match.",
        "type": "Functional", "priority": "Medium",
    },
    {
        "id": "TC_008", "module": "Customer",
        "title": "Add duplicate customer with same phone number",
        "pre": "Customer with phone 9876543210 already exists.",
        "steps": "1. Navigate to Customers → Add Customer.\n2. Enter all fields; set phone to 9876543210.\n3. Click Save.",
        "expected": "HTTP 400/409; error message 'Customer with this phone already exists'; no duplicate record created.",
        "type": "Negative", "priority": "High",
    },
    {
        "id": "TC_009", "module": "Customer",
        "title": "Update customer email address",
        "pre": "Customer record with a valid ID exists.",
        "steps": "1. Navigate to Customers.\n2. Open an existing customer.\n3. Change email to newemail@test.com.\n4. Click Update.",
        "expected": "HTTP 200; customer record reflects updated email; audit log captures the change event.",
        "type": "Functional", "priority": "Medium",
    },
    {
        "id": "TC_010", "module": "Customer",
        "title": "Add customer with invalid email format",
        "pre": "Logged in as SALES_EXECUTIVE.",
        "steps": "1. Navigate to Customers → Add Customer.\n2. Enter email as 'notanemail'.\n3. Fill remaining fields correctly.\n4. Click Save.",
        "expected": "Form validation rejects input; error shown next to email field; API call NOT made.",
        "type": "Negative", "priority": "Medium",
    },
    # ── Booking ──
    {
        "id": "TC_011", "module": "Booking",
        "title": "Create booking with valid customer and vehicle",
        "pre": "Customer and vehicle with status IN_SHOWROOM exist; logged in as SALES_EXECUTIVE.",
        "steps": "1. Navigate to Bookings → New Booking.\n2. Select existing customer.\n3. Select available vehicle.\n4. Set booking date and expected delivery date.\n5. Choose payment mode.\n6. Click Create Booking.",
        "expected": "HTTP 201; booking created with status PENDING; vehicle status changes to BOOKED; booking appears in list.",
        "type": "Functional", "priority": "High",
    },
    {
        "id": "TC_012", "module": "Booking",
        "title": "Create booking with missing mandatory fields",
        "pre": "Logged in as SALES_EXECUTIVE.",
        "steps": "1. Navigate to Bookings → New Booking.\n2. Leave customer field blank.\n3. Select a vehicle.\n4. Click Create Booking.",
        "expected": "HTTP 400 / validation error; error indicates missing required fields; booking NOT created.",
        "type": "Negative", "priority": "High",
    },
    {
        "id": "TC_013", "module": "Booking",
        "title": "Cancel an existing PENDING booking",
        "pre": "A booking with status PENDING exists.",
        "steps": "1. Navigate to Bookings.\n2. Open target booking.\n3. Click Cancel Booking.\n4. Confirm cancellation.",
        "expected": "HTTP 200; booking status updates to CANCELLED; vehicle availability reverts; cancellation reflected in list.",
        "type": "Functional", "priority": "High",
    },
    {
        "id": "TC_014", "module": "Booking",
        "title": "Create booking with non-existing customer ID",
        "pre": "Logged in as SALES_EXECUTIVE; backend accessible.",
        "steps": "1. Send POST /api/bookings with payload containing non-existent customerId e.g. 999999.",
        "expected": "HTTP 404; error message 'Customer not found'; booking record NOT created.",
        "type": "Negative", "priority": "High",
    },
    {
        "id": "TC_015", "module": "Booking",
        "title": "Filter bookings by status and date range",
        "pre": "Multiple bookings exist in different statuses.",
        "steps": "1. Navigate to Bookings.\n2. Apply filter: status = CONFIRMED, date range = current month.\n3. Observe results.",
        "expected": "Only CONFIRMED bookings within selected date range displayed; pagination count updated accordingly.",
        "type": "Functional", "priority": "Medium",
    },
    # ── Inventory ──
    {
        "id": "TC_016", "module": "Inventory",
        "title": "Add new vehicle with complete details",
        "pre": "Logged in as DEALER_MANAGER; dealership and variant exist.",
        "steps": "1. Navigate to Inventory → Add Vehicle.\n2. Fill VIN, model, variant, color, year, fuel type, transmission, selling price.\n3. Click Save.",
        "expected": "HTTP 201; vehicle added with status IN_SHOWROOM; record visible in inventory list with all entered details.",
        "type": "Functional", "priority": "High",
    },
    {
        "id": "TC_017", "module": "Inventory",
        "title": "Vehicle availability validation before booking",
        "pre": "Vehicle with status BOOKED exists.",
        "steps": "1. Navigate to Bookings → New Booking.\n2. Attempt to select the already-BOOKED vehicle.",
        "expected": "BOOKED vehicle is excluded from available vehicles dropdown or shows a warning; booking cannot be finalised.",
        "type": "Functional", "priority": "High",
    },
    {
        "id": "TC_018", "module": "Inventory",
        "title": "Update vehicle status to SOLD after delivery",
        "pre": "Vehicle with status DISPATCHED exists.",
        "steps": "1. Navigate to Inventory.\n2. Select the dispatched vehicle.\n3. Update status to SOLD.",
        "expected": "HTTP 200; vehicle status set to SOLD; associated booking status updated to COMPLETED.",
        "type": "Functional", "priority": "Medium",
    },
    # ── Admin & Security ──
    {
        "id": "TC_019", "module": "Admin",
        "title": "SALES_EXECUTIVE cannot access Admin Settings",
        "pre": "Logged in as SALES_EXECUTIVE.",
        "steps": "1. Log in as SALES_EXECUTIVE.\n2. Navigate to /admin-settings.\n3. Alternatively send GET /api/users with the SALES_EXECUTIVE JWT token.",
        "expected": "HTTP 403 Forbidden; user redirected to /unauthorized page; no user management data exposed.",
        "type": "Security", "priority": "High",
    },
    {
        "id": "TC_020", "module": "Admin",
        "title": "Add new user and assign role",
        "pre": "Logged in as SUPER_ADMIN; Admin Settings page accessible.",
        "steps": "1. Navigate to Admin Settings → Users.\n2. Click Add User.\n3. Enter full name, email, password, phone and role SALES_EXECUTIVE.\n4. Click Save.",
        "expected": "HTTP 201; new user listed; can log in with assigned credentials; role persisted correctly in DB.",
        "type": "Functional", "priority": "High",
    },
    # ── Edge / Boundary ──
    {
        "id": "TC_021", "module": "Customer",
        "title": "Add customer with excessively long name (boundary)",
        "pre": "Logged in as SALES_EXECUTIVE.",
        "steps": "1. Navigate to Customers → Add Customer.\n2. Enter first name with 300+ characters.\n3. Fill other fields normally.\n4. Click Save.",
        "expected": "HTTP 400; error indicating max field length exceeded; record NOT created; no DB exception shown on frontend.",
        "type": "Boundary", "priority": "Medium",
    },
    {
        "id": "TC_022", "module": "Booking",
        "title": "Booking amount set to zero or negative value",
        "pre": "Logged in as SALES_EXECUTIVE.",
        "steps": "1. Navigate to Bookings → New Booking.\n2. Enter all valid fields but set booking amount to 0 or -500.\n3. Click Create Booking.",
        "expected": "HTTP 400; validation error 'Booking amount must be greater than 0'; booking NOT saved.",
        "type": "Boundary", "priority": "Medium",
    },
    # ── Reports & Audit ──
    {
        "id": "TC_023", "module": "Reports",
        "title": "Audit log captures user action",
        "pre": "Logged in as SUPER_ADMIN; audit logging enabled.",
        "steps": "1. Navigate to Admin Settings → Users.\n2. Toggle active status of an existing user.\n3. Navigate to Reports → Audit Logs.",
        "expected": "Audit log contains entry with action TOGGLE_USER_STATUS, affected user ID, new status, and timestamp.",
        "type": "Functional", "priority": "Medium",
    },
]

# ──────────────────────────────────────────────
# Group test cases by module
# ──────────────────────────────────────────────
from collections import OrderedDict

MODULE_ORDER = [
    "Authentication", "Customer", "Booking", "Inventory", "Admin", "Reports"
]

def group_by_module(tcs):
    groups = OrderedDict()
    for m in MODULE_ORDER:
        groups[m] = [tc for tc in tcs if tc["module"] == m]
    return groups

# ──────────────────────────────────────────────
# Build document
# ──────────────────────────────────────────────
def build_doc():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)
        section.page_width    = Cm(33.87)  # A3 landscape width
        section.page_height   = Cm(21.0)   # A3 height
        section.orientation   = 1          # landscape

    # ── Cover block ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Hyundai Dealer Management System")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = HYUNDAI_BLUE

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Quality Assurance – Test Case Document")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = ACCENT_BLUE
    sub_run.bold = True

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"Version 1.0  ·  Prepared by QA Team  ·  Date: {datetime.date.today().strftime('%d %B %Y')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    meta_run.italic = True

    doc.add_paragraph()  # spacer

    # ── Divider line (table with 1 row & 1 col, navy bg) ──
    div = doc.add_table(rows=1, cols=1)
    div.style = "Table Grid"
    div_cell = div.cell(0, 0)
    set_cell_bg(div_cell, HYUNDAI_BLUE)
    div_cell.paragraphs[0].add_run(" ")
    div.rows[0].height = Cm(0.08)

    doc.add_paragraph()  # spacer

    # ── Project meta table ──
    mt = doc.add_table(rows=5, cols=4)
    mt.style = "Table Grid"
    meta_items = [
        ("Project",     "Hyundai DMS",    "Backend",     "Spring Boot (Java)"),
        ("Frontend",    "React + Vite",   "Database",    "MySQL"),
        ("Auth",        "JWT Bearer Token","Roles",      "SUPER_ADMIN · DEALER_MANAGER · SALES_EXECUTIVE · SENIOR_OFFICIAL"),
        ("Total TCs",   "23",             "Prepared By", "QA Team"),
        ("Sprint",      "Current Sprint", "Status",      "Ready for Execution"),
    ]
    for r_idx, (k1, v1, k2, v2) in enumerate(meta_items):
        row = mt.rows[r_idx]
        for c_idx, (text, is_key) in enumerate([(k1,True),(v1,False),(k2,True),(v2,False)]):
            cell = row.cells[c_idx]
            if is_key:
                set_cell_bg(cell, RGBColor(0xE3, 0xEA, 0xF8))
            cell_para(cell, text, bold=is_key, size=9,
                      color=HYUNDAI_BLUE if is_key else DARK_TEXT)

    doc.add_paragraph()

    # ── Section legend ──
    leg_para = doc.add_paragraph()
    leg_run = leg_para.add_run(
        "LEGEND   |   Test Types: Functional  ·  Negative  ·  Security  ·  Boundary   |   "
        "Priority: High  ·  Medium  ·  Low"
    )
    leg_run.font.size = Pt(8.5)
    leg_run.italic = True
    leg_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ── Section tables ──
    grouped = group_by_module(TEST_CASES)

    for module_name, tcs in grouped.items():
        if not tcs:
            continue

        # Section heading
        sec_heading = doc.add_paragraph()
        sec_run = sec_heading.add_run(f"  {module_name.upper()} MODULE  ")
        sec_run.bold = True
        sec_run.font.size = Pt(12)
        sec_run.font.color.rgb = WHITE
        # shade via highlight workaround: use a 1-cell table
        sec_tbl = doc.add_table(rows=1, cols=1)
        sec_tbl.style = "Table Grid"
        sec_cell = sec_tbl.cell(0, 0)
        set_cell_bg(sec_cell, HYUNDAI_BLUE)
        p = sec_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f"  {module_name.upper()} MODULE  —  {len(tcs)} Test Case(s)")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = WHITE
        sec_cell.height = Cm(0.8)

        # Data table
        table = doc.add_table(rows=1 + len(tcs), cols=len(COLUMNS))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_col_widths(table)
        add_table_header(table)

        for row_i, tc in enumerate(tcs):
            actual_row = table.rows[row_i + 1]
            is_alt = row_i % 2 == 1
            row_bg = ALT_ROW if is_alt else WHITE

            values = [
                tc["id"], tc["module"], tc["title"],
                tc["pre"], tc["steps"], tc["expected"],
                tc["type"], tc["priority"],
            ]

            for col_i, val in enumerate(values):
                cell = actual_row.cells[col_i]
                set_cell_border(cell, "B0C4DE")
                set_cell_bg(cell, row_bg)

                # Special styling for TC ID
                if col_i == 0:
                    cell_para(cell, val, bold=True, color=ACCENT_BLUE, size=8.5,
                              align=WD_ALIGN_PARAGRAPH.CENTER)
                # Special colouring for Test Type
                elif col_i == 6:
                    tc_color = TYPE_COLORS.get(val, DARK_TEXT)
                    cell_para(cell, val, bold=True, color=tc_color, size=8.5,
                              align=WD_ALIGN_PARAGRAPH.CENTER)
                # Special colouring for Priority
                elif col_i == 7:
                    pr_color = PRIORITY_COLORS.get(val, DARK_TEXT)
                    cell_para(cell, val, bold=True, color=pr_color, size=8.5,
                              align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    cell_para(cell, val, size=8.5)

        doc.add_paragraph()  # spacer between modules

    # ── Footer note ──
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_para.add_run(
        "This document is auto-generated by the QA pipeline  ·  "
        "Hyundai DMS · Confidential  ·  "
        f"Generated: {datetime.datetime.now().strftime('%d %b %Y %H:%M')}"
    )
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    out_path = "DMS_Test_Cases.docx"
    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")

if __name__ == "__main__":
    build_doc()
